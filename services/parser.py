import fitz  # PyMuPDF
import re
import os
from docx import Document

SECTION_HEADER_ALIASES = {
    "summary": "summary",
    "professional summary": "summary",
    "profile": "summary",
    "objective": "summary",
    "skills": "skills",
    "technical skills": "skills",
    "core skills": "skills",
    "technologies": "skills",
    "experience": "experience",
    "work experience": "experience",
    "professional experience": "experience",
    "employment history": "experience",
    "projects": "projects",
    "project experience": "projects",
    "education": "education",
    "academic background": "education",
    "certifications": "certifications",
    "certificates": "certifications",
}

SKILL_SIGNAL_MAP = {
    "python": r"\bpython\b",
    "java": r"\bjava\b",
    "javascript": r"\bjavascript\b",
    "typescript": r"\btypescript\b",
    "react": r"\breact\b",
    "node": r"\bnode(?:\.js|js)?\b",
    "express": r"\bexpress(?:\.js|js)?\b",
    "spring": r"\bspring(?:\s+boot)?\b",
    "django": r"\bdjango\b",
    "flask": r"\bflask\b",
    "fastapi": r"\bfastapi\b",
    "sql": r"\bsql\b",
    "postgresql": r"\bpostgresql\b",
    "mysql": r"\bmysql\b",
    "mongodb": r"\bmongodb\b",
    "redis": r"\bredis\b",
    "kafka": r"\bkafka\b",
    "docker": r"\bdocker\b",
    "kubernetes": r"\bkubernetes\b",
    "terraform": r"\bterraform\b",
    "jenkins": r"\bjenkins\b",
    "git": r"\bgit\b",
    "aws": r"\baws\b",
    "azure": r"\bazure\b",
    "gcp": r"\bgcp\b",
    "ci/cd": r"\bci/cd\b",
    "power bi": r"\bpower\s*bi\b",
    "tableau": r"\btableau\b",
    "scikit-learn": r"\bscikit-?learn\b",
    "pandas": r"\bpandas\b",
    "numpy": r"\bnumpy\b",
    "tensorflow": r"\btensorflow\b",
    "pytorch": r"\bpytorch\b",
    "figma": r"\bfigma\b",
    "selenium": r"\bselenium\b",
    "kotlin": r"\bkotlin\b",
    "swift": r"\bswift\b",
    "swiftui": r"\bswiftui\b",
    "uikit": r"\buikit\b",
    "android": r"\bandroid\b",
    "ruby": r"\bruby\b",
    "rails": r"\brails\b",
    "rspec": r"\brspec\b",
    "php": r"\bphp\b",
    "laravel": r"\blaravel\b",
    "symfony": r"\bsymfony\b",
    "go": r"\bgo(lang)?\b",
    "grpc": r"\bgrpc\b",
    "gin": r"\bgin\b",
    "gorm": r"\bgorm\b",
    "c++": r"\bc\+\+\b|\bcpp\b",
    "c": r"\bc\b",
    "dotnet": r"\b\.net\b|\bdotnet\b",
    "asp.net": r"\basp\.?net\b",
    "entity framework": r"\bentity framework\b|\bef core\b",
}

HEADER_HINTS = {
    "summary": {"summary", "profile", "objective"},
    "skills": {"skills", "technologies", "tech stack", "core competencies", "competencies"},
    "experience": {"experience", "employment", "work history", "professional experience"},
    "projects": {"projects", "project experience", "case studies"},
    "education": {"education", "academic", "qualification"},
    "certifications": {"certifications", "certificates", "licenses"},
}

CERTIFICATION_KEYWORDS = [
    "aws certified", "azure certified", "google cloud", "gcp", "cka", "ckad", "terraform associate",
    "pmp", "scrum master", "istqb", "comptia", "security+", "cissp", "ceh",
]


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\u2013\u2014\u2212]", "-", text)
    text = re.sub(r"[\u2022\u25CF\u25E6\u00B7]", "•", text)

    normalized_lines = []
    for line in text.splitlines():
        cleaned = re.sub(r"\s+", " ", line).strip()
        if cleaned:
            normalized_lines.append(cleaned)

    return "\n".join(normalized_lines)

def extract_text_from_pdf(filepath: str) -> str:
    doc = fitz.open(filepath)
    parts = []
    for page in doc:
        parts.append(page.get_text("text"))
    doc.close()
    return "\n".join(parts)


def extract_text_from_docx(filepath: str) -> str:
    doc = Document(filepath)
    paragraphs = []

    for p in doc.paragraphs:
        if not p.text or not p.text.strip():
            continue
        line = p.text.strip()
        style_name = (p.style.name or "").lower() if p.style else ""
        if "list" in style_name and not re.match(r"^[-•*]", line):
            line = f"• {line}"
        paragraphs.append(line)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _normalize_header_candidate(line: str) -> str:
    line = line.strip().lower()
    line = re.sub(r"^[^a-z0-9]+|[^a-z0-9]+$", "", line)
    line = re.sub(r"[:\-]+$", "", line).strip()
    line = re.sub(r"\s+", " ", line)
    return line


def _detect_section_header(line: str) -> str:
    candidate = _normalize_header_candidate(line)
    if candidate in SECTION_HEADER_ALIASES:
        return SECTION_HEADER_ALIASES[candidate]

    for section, hints in HEADER_HINTS.items():
        if candidate in hints:
            return section
        if any(hint in candidate for hint in hints):
            return section

    return ""


def _looks_like_header(line: str) -> bool:
    clean = line.strip()
    if not clean:
        return False

    alnum = re.sub(r"[^A-Za-z0-9 ]", "", clean)
    words = [w for w in alnum.split() if w]
    if not words or len(words) > 6:
        return False

    if _detect_section_header(clean):
        return True

    if clean.isupper() and 1 <= len(words) <= 4:
        return True

    # Title-like short line is a likely heading.
    title_like = sum(1 for w in words if w[:1].isupper())
    if title_like >= max(1, len(words) - 1) and len(words) <= 4:
        return True

    return False


def _infer_section_from_line(line: str) -> str:
    l = line.lower()
    if re.search(r"\b(b\.?(tech|e)|m\.?(tech|e|s|ba)|bsc|msc|bachelor|master|university|college|cgpa|gpa)\b", l):
        return "education"
    if re.search(r"\b(intern|engineer|developer|manager|analyst|architect|consultant)\b", l):
        return "experience"
    if re.search(r"\b(project|built|developed|implemented|deployed|designed)\b", l):
        return "projects"
    if re.search(r"\b(certified|certification|aws certified|azure certified|scrum|pmp)\b", l):
        return "certifications"
    return ""

def split_sections(text: str) -> dict:
    lines = text.splitlines()
    sections = {}
    current = "other"
    sections[current] = []

    for line in lines:
        clean = line.strip()
        if not clean:
            continue

        detected_header = _detect_section_header(clean)
        if detected_header:
            current = detected_header
            sections.setdefault(current, [])
            continue

        if _looks_like_header(clean):
            possible = _detect_section_header(clean)
            if possible:
                current = possible
                sections.setdefault(current, [])
                continue

        if current == "other":
            inferred = _infer_section_from_line(clean)
            if inferred:
                current = inferred
                sections.setdefault(current, [])

        sections.setdefault(current, []).append(clean)

    final_sections = {}
    for key, value in sections.items():
        final_sections[key] = "\n".join(value).strip()

    return final_sections


def extract_name(text: str) -> str:
    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            continue

        if re.search(r"@|https?://|linkedin|github|portfolio", clean.lower()):
            continue

        if re.search(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", clean):
            continue

        if _detect_section_header(clean):
            continue

        if 2 <= len(clean.split()) <= 6:
            return clean
    return ""


def extract_skills(skills_text: str) -> list:
    if not skills_text:
        return []

    raw = re.split(r"[,•;|/\n]+", skills_text)
    cleaned = set()

    for token in raw:
        skill = token.strip().lower()
        skill = re.sub(r"\s+", " ", skill)
        skill = re.sub(r"^[\-:]+|[\-:]+$", "", skill).strip()
        if len(skill) < 2:
            continue
        cleaned.add(skill)

    return sorted(cleaned)


def extract_skills_from_full_text(full_text: str) -> list:
    normalized = full_text.lower()
    found = set()
    for skill, pattern in SKILL_SIGNAL_MAP.items():
        if re.search(pattern, normalized):
            found.add(skill)
    return sorted(found)


def extract_profile_signals(full_text: str, sections: dict) -> dict:
    text = full_text.lower()
    cert_section = sections.get("certifications", "").lower()
    projects_section = sections.get("projects", "").lower()
    experience_section = sections.get("experience", "").lower()

    cert_hits = 0
    for keyword in CERTIFICATION_KEYWORDS:
        cert_hits += len(re.findall(re.escape(keyword), text))

    # Links found in project/experience sections usually indicate portfolio/project evidence.
    project_links = len(re.findall(r"https?://", projects_section)) + len(re.findall(r"github\.com", projects_section))
    experience_links = len(re.findall(r"https?://", experience_section)) + len(re.findall(r"github\.com", experience_section))

    quantified_impact = bool(re.search(r"\b\d+(?:\.\d+)?%\b|\b(increased|reduced|improved)\b", text))

    return {
        "github": bool(re.search(r"github\.com|\bgithub\b", text)),
        "linkedin": bool(re.search(r"linkedin\.com|\blinkedin\b", text)),
        "portfolio": bool(re.search(r"\bportfolio\b|behance\.net|dribbble\.com|medium\.com|personal website", text)),
        "kaggle": bool(re.search(r"kaggle\.com|\bkaggle\b", text)),
        "huggingface": bool(re.search(r"huggingface\.co|\bhugging face\b", text)),
        "figma_link": bool(re.search(r"figma\.com|\bfigma\b", text)),
        "behance": bool(re.search(r"behance\.net|\bbehance\b", text)),
        "dribbble": bool(re.search(r"dribbble\.com|\bdribbble\b", text)),
        "certifications_count": cert_hits + (1 if cert_section.strip() else 0),
        "project_links_count": project_links,
        "experience_links_count": experience_links,
        "project_mentions": len(re.findall(r"\b(project|case study|product)\b", projects_section or text)),
        "has_quantified_impact": quantified_impact,
    }

def parse_resume(filepath: str) -> dict:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(filepath)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(filepath)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    full_text = normalize_text(raw_text)

    sections = split_sections(full_text)

    name = extract_name(full_text)
    skills_sources = [
        sections.get("skills", ""),
        sections.get("summary", ""),
        sections.get("experience", ""),
        sections.get("projects", ""),
    ]
    merged_skill_text = "\n".join([s for s in skills_sources if s])
    skills = set(extract_skills(merged_skill_text))
    skills.update(extract_skills_from_full_text(full_text))

    if not sections.get("skills") and skills:
        sections["skills"] = ", ".join(sorted(skills))

    profile_signals = extract_profile_signals(full_text, sections)

    return {
        "name": name,
        "full_text": full_text,
        "sections": sections,
        "skills": sorted(skills),
        "profile_signals": profile_signals,
    }
